"""
Document service for file processing and content extraction
"""
import os
import uuid
from typing import Optional, Dict, Any
import pandas as pd
from docx import Document as DocxDocument
import PyPDF2
from io import BytesIO
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException, status
from app.models.document import Document, DocumentType, DocumentStatus
from app.core.config import settings
from app.workflows.adjustment_workflow import adjustment_workflow
import aiofiles

class DocumentService:
    def __init__(self, db: Session):
        self.db = db
        self.upload_dir = settings.UPLOAD_DIR
        os.makedirs(self.upload_dir, exist_ok=True)
    
    async def upload_document(self, project_id: int, file: UploadFile) -> Document:
        """Upload and process a document"""
        # Validate file type
        allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "text/csv"]
        
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unsupported file type"
            )
        
        # Check file size
        file_content = await file.read()
        if len(file_content) > settings.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File size exceeds maximum allowed size"
            )
        
        # Generate unique filename
        file_extension = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = os.path.join(self.upload_dir, unique_filename)
        
        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Create document record
        document = Document(
            project_id=project_id,
            filename=unique_filename,
            original_filename=file.filename,
            file_path=file_path,
            file_size=len(file_content),
            mime_type=file.content_type,
            status=DocumentStatus.PENDING
        )
        
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        
        # Process document asynchronously
        await self._process_document(document)
        
        return document
    
    async def _process_document(self, document: Document):
        """Process document content extraction and classification"""
        try:
            # Extract content based on file type
            content = await self._extract_content(document.file_path, document.mime_type)
            
            # Classify document type
            doc_type = await self._classify_document(content, document.original_filename)
            
            # Update document with extracted content
            document.raw_text = content[:10000]  # Store first 10k characters
            document.document_type = doc_type
            document.status = DocumentStatus.PROCESSED
            
            self.db.commit()
            
            # Trigger adjustment analysis
            await self._analyze_for_adjustments(document)
            
        except Exception as e:
            document.status = DocumentStatus.FAILED
            document.processing_error = str(e)
            self.db.commit()
    
    async def _extract_content(self, file_path: str, mime_type: str) -> str:
        """Extract text content from various file types"""
        if mime_type == "application/pdf":
            return await self._extract_pdf_content(file_path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await self._extract_docx_content(file_path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return await self._extract_excel_content(file_path)
        elif mime_type == "text/csv":
            return await self._extract_csv_content(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    async def _extract_excel_content(self, file_path: str) -> str:
        """Extract data from Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                content += f"Sheet: {sheet_name}\n"
                content += df.to_string() + "\n\n"
            
            return content
        except Exception as e:
            raise ValueError(f"Error reading Excel file: {str(e)}")
    
    async def _extract_csv_content(self, file_path: str) -> str:
        """Extract data from CSV file"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            raise ValueError(f"Error reading CSV file: {str(e)}")
    
    async def _classify_document(self, content: str, filename: str) -> DocumentType:
        """Classify document type based on content and filename"""
        filename_lower = filename.lower()
        content_lower = content.lower()
        
        # Simple classification based on keywords
        if any(keyword in filename_lower for keyword in ["gl", "general", "ledger"]):
            return DocumentType.GL
        elif any(keyword in filename_lower for keyword in ["p&l", "profit", "loss", "income"]):
            return DocumentType.PL
        elif any(keyword in filename_lower for keyword in ["payroll", "salary", "wages"]):
            return DocumentType.PAYROLL
        elif any(keyword in filename_lower for keyword in ["trial", "balance"]):
            return DocumentType.TRIAL_BALANCE
        elif any(keyword in content_lower for keyword in ["general ledger", "account", "debit", "credit"]):
            return DocumentType.GL
        elif any(keyword in content_lower for keyword in ["revenue", "expense", "profit", "loss", "ebitda"]):
            return DocumentType.PL
        else:
            return DocumentType.OTHER
    
    async def _analyze_for_adjustments(self, document: Document):
        """Analyze document for potential adjustments using LangGraph workflow"""
        try:
            # Get project context
            project = document.project
            project_context = {
                "project_name": project.name,
                "client_name": project.client_name,
                "materiality_amount": project.materiality_amount,
                "materiality_percentage": project.materiality_percentage
            }
            
            # Prepare workflow state
            workflow_state = {
                "document_content": document.raw_text,
                "document_type": document.document_type.value,
                "project_context": project_context,
                "identified_adjustments": [],
                "processed_adjustments": [],
                "materiality_threshold": project.materiality_amount,
                "materiality_percentage": project.materiality_percentage
            }
            
            # Run the adjustment workflow
            result = await adjustment_workflow.process_document(workflow_state)
            
            # Store the analysis results
            document.extracted_data = {
                "adjustments_identified": len(result.get("processed_adjustments", [])),
                "analysis_completed": True,
                "workflow_result": result
            }
            
            self.db.commit()
            
        except Exception as e:
            print(f"Error in adjustment analysis: {str(e)}")
            # Don't fail the entire document processing if adjustment analysis fails
    
    async def get_document(self, document_id: int) -> Optional[Document]:
        """Get document by ID"""
        return self.db.query(Document).filter(Document.id == document_id).first()
    
    async def get_project_documents(self, project_id: int) -> list[Document]:
        """Get all documents for a project"""
        return self.db.query(Document).filter(Document.project_id == project_id).all()
    
    async def delete_document(self, document_id: int) -> bool:
        """Delete a document"""
        document = await self.get_document(document_id)
        if not document:
            return False
        
        # Delete file from filesystem
        try:
            os.remove(document.file_path)
        except FileNotFoundError:
            pass
        
        # Delete from database
        self.db.delete(document)
        self.db.commit()
        
        return True